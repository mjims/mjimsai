"""
Knowledge base routes — upload and manage training documents.
"""

from __future__ import annotations

import uuid

from fastapi import APIRouter, Depends, File, HTTPException, UploadFile, status
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.api.deps import get_current_user
from app.config import get_settings
from app.database import get_db
from app.models.knowledge import KnowledgeDocument, KnowledgeChunk
from app.models.user import User
from app.services import agent_service

router = APIRouter(prefix="/agents/{agent_id}/knowledge", tags=["Knowledge Base"])


async def _extract_text(file: UploadFile, content: bytes) -> str:
    filename = file.filename or ""
    content_type = file.content_type or ""

    if content_type == "text/plain" or filename.endswith(".txt"):
        return content.decode("utf-8", errors="replace")
    if content_type == "text/markdown" or filename.endswith(".md"):
        return content.decode("utf-8", errors="replace")
    if content_type == "application/pdf" or filename.endswith(".pdf"):
        try:
            from PyPDF2 import PdfReader
            import io
            reader = PdfReader(io.BytesIO(content))
            return "\n\n".join(p.extract_text() or "" for p in reader.pages)
        except Exception as e:
            raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail=f"PDF error: {e}")
    if filename.endswith(".docx"):
        try:
            from docx import Document
            import io
            doc = Document(io.BytesIO(content))
            return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as e:
            raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail=f"DOCX error: {e}")

    raise HTTPException(status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
                        detail=f"Unsupported type: {content_type}. Supported: .txt, .md, .pdf, .docx")


def _chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> list[str]:
    if len(text) <= chunk_size:
        return [text]
    chunks, start = [], 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        if end < len(text):
            bp = max(chunk.rfind("."), chunk.rfind("\n"))
            if bp > chunk_size * 0.5:
                chunk = text[start:start + bp + 1]
                end = start + bp + 1
        chunks.append(chunk.strip())
        start = end - overlap
    return [c for c in chunks if c]


@router.post("", status_code=status.HTTP_201_CREATED)
async def upload_document(
    agent_id: uuid.UUID,
    file: UploadFile = File(...),
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    settings = get_settings()
    agent = await agent_service.get_agent_by_id(db, user.id, agent_id)
    if not agent:
        raise HTTPException(status_code=404, detail="Agent not found")

    content = await file.read()
    if len(content) > settings.MAX_UPLOAD_SIZE_MB * 1024 * 1024:
        raise HTTPException(status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
                            detail=f"File too large. Max: {settings.MAX_UPLOAD_SIZE_MB}MB")

    extracted_text = await _extract_text(file, content)
    if not extracted_text.strip():
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="No text extracted")

    doc = KnowledgeDocument(
        agent_id=agent.id, filename=file.filename or "unknown",
        content_type=file.content_type or "application/octet-stream",
        file_size_bytes=len(content), content_text=extracted_text, status="ready",
    )
    db.add(doc)
    await db.flush()

    chunks = _chunk_text(extracted_text)
    for idx, chunk_text in enumerate(chunks):
        db.add(KnowledgeChunk(
            document_id=doc.id, content=chunk_text,
            chunk_index=idx, token_count=len(chunk_text) // 4,
        ))
    doc.chunk_count = len(chunks)
    await db.commit()

    return {"id": str(doc.id), "filename": doc.filename, "content_type": doc.content_type,
            "file_size_bytes": doc.file_size_bytes, "chunk_count": doc.chunk_count, "status": doc.status}


@router.get("")
async def list_documents(
    agent_id: uuid.UUID,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    agent = await agent_service.get_agent_by_id(db, user.id, agent_id)
    if not agent:
        raise HTTPException(status_code=404, detail="Agent not found")

    result = await db.execute(
        select(KnowledgeDocument).where(KnowledgeDocument.agent_id == agent.id)
        .order_by(KnowledgeDocument.created_at.desc())
    )
    docs = result.scalars().all()
    return {"documents": [{"id": str(d.id), "filename": d.filename, "content_type": d.content_type,
                           "file_size_bytes": d.file_size_bytes, "chunk_count": d.chunk_count,
                           "status": d.status, "created_at": d.created_at.isoformat()} for d in docs]}


@router.delete("/{document_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_document(
    agent_id: uuid.UUID,
    document_id: uuid.UUID,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    agent = await agent_service.get_agent_by_id(db, user.id, agent_id)
    if not agent:
        raise HTTPException(status_code=404, detail="Agent not found")

    result = await db.execute(
        select(KnowledgeDocument).where(
            KnowledgeDocument.id == document_id,
            KnowledgeDocument.agent_id == agent.id,
        )
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    await db.delete(doc)
    await db.commit()
